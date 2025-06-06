#Importing modules
import pdfplumber
import docx
#Last updated - 11:33 a.m. - 11-04-2025
#Assigning letters
# Define the dictionary to verify
dict0 = {
    'A': 'M', 'B': 'Q', 'C': 'W', 'D': 'E', 'E': 'R', 'F': 'T', 'G': 'Y', 'H': 'U', 'I': 'I', 'J': 'O', 'K': 'P', 'L': 'A',
    'M': 'S', 'N': 'D', 'O': 'F', 'P': 'G', 'Q': 'H', 'R': 'J', 'S': 'K', 'T': 'L', 'U': 'Z', 'V': 'X', 'W': 'C', 'X': 'V',
    'Y': 'B', 'Z': 'N', 
    'a': 'm', 'b': 'q', 'c': 'w', 'd': 'e', 'e': 'r', 'f': 't', 'g': 'y', 'h': 'u', 'i': 'i', 'j': 'o', 'k': 'p', 'l': 'a',
    'm': 's', 'n': 'd', 'o': 'f', 'p': 'g', 'q': 'h', 'r': 'j', 's': 'k', 't': 'l', 'u': 'z', 'v': 'x', 'w': 'c', 'x': 'v',
    'y': 'b', 'z': 'n',
    '0': '9', '1': '8', '2': '7', '3': '6', '4': '5', '5': '4', '6': '3', '7': '2', '8': '1', '9': '0',
    ' ': '#', '#': ' ',  # Added this pair to ensure symmetry
    '.': ',', ',': '.', '!': '?', '?': '!', '-': '_', '_': '-', '(': ')', ')': '(', '[': ']', ']': '[', 
    '{': '}', '}': '{', '<': '>', '>': '<', '@': '$', '$': '@', '%': '^', '^': '%', '&': '*', '*': '&', '+': '=', 
    '=': '+', '/': '\\', '\\': '/', '|': ';', ';': '|', ':': "'", "'": ':', '"': '`', '`': '"'
}


#making its reverse dictionary

reverse_dict0 = dict(zip(dict0.values(), dict0.keys()))

# Usual Encoder func.

def encoder_usual(x00):
    x1 = x00.upper()
    ux1 = list(x1)
    for n in range(len(ux1)):
        something = ux1[n]
        ux1[n] = dict0.get(something, something)
    encoded_message = ''.join(ux1)
    print("Your encoded message is - ",encoded_message, "\n")
    
# Txt Encoder Function

def encoder_txt(filename, savingfilename):
    if filename.endswith(".txt") == False:
        finalfile = filename + ".txt"
    else:
        finalfile = filename
        
    file = open(finalfile, "r") 
    intext = file.read()
    ux1 = list(intext)
    for n in range(len(ux1)):
        something = ux1[n]
        ux1[n] = dict0.get(something, something)
        
    encoded_message = ''.join(ux1)
    newfile = open(savingfilename, "w")
    newfile.write(encoded_message)
    newfile.close()
    file.close()
    
    
#Pdf encoding function
def encoder_pdf(filename, saving_filename00):
    if saving_filename00.endswith('.txt') == False:
        saving_filename01 = saving_filename00 + '.txt'
    else:
         saving_filename01 = saving_filename00
    pdf = pdfplumber.open(filename)
    pdf_text = ""
    for page in pdf.pages:
        if page.extract_text() == None:
            pdf_text += "No text found for this page" + "\n-------Next page------\n"
        else:
            pdf_text += page.extract_text() + "\n-------Next page------\n"
    pdf.close()    
    ux1 = list(pdf_text)
    for n in range(len(ux1)):
        something = ux1[n]
        ux1[n] = dict0.get(something, something)
    encoded_message = ''.join(ux1)
    savingfile = open(saving_filename01, "w")
    savingfile.write(encoded_message)
    savingfile.close()
    print(f"\nCongrats!! File saved as {saving_filename01}\n")

    
#Docx encoder function :-)

def encoder_docx(filename, saving_filename00):
    if filename.endswith(".docx") == False:
        finalfilename = filename + ".docx"
    else:
        finalfilename = filename
    
    
    if saving_filename00.endswith('.txt') == False:
        saving_filename01 = saving_filename00 + '.txt'
    else:
         saving_filename01 = saving_filename00
         
    doc = docx.Document(finalfilename)
    doc_text = ""
    for para in doc.paragraphs:
        if para.text == "":
            doc_text += "No text found for this para" + "\n-------Next para------\n"
        else:
            doc_text += para.text + "\n-------Next page------\n"
   
    ux1 = list(doc_text)
    for n in range(len(ux1)):
        something = ux1[n]
        ux1[n] = dict0.get(something, something)
    encoded_message = ''.join(ux1)
    savingfile = open(saving_filename01, "w", encoding="utf-8")
    savingfile.write(encoded_message)
    savingfile.close()
    print(f"\nCongrats!! File saved as {saving_filename01}\n") 
                
#Decoder func.

def decoder_usual(message):
    ux1 = list(message)
    for n in range(len(ux1)):
        something = ux1[n]
        ux1[n] = reverse_dict0.get(something, something)
    decoded_message = ''.join(ux1)
    print("Your decoded message is - ", decoded_message, "\n")
    
#Txt Decoder function

def decoder_txt(filename, savingfilename):
    if filename.endswith(".txt") == False:
        finalfile = filename + ".txt"
    else:
        finalfile = filename
        
    file = open(finalfile, "r") 
    intext = file.read()
    ux1 = list(intext)
    for n in range(len(ux1)):
        something = ux1[n]
        ux1[n] = reverse_dict0.get(something, something)
        
    encoded_message = ''.join(ux1)
    newfile = open(savingfilename, "w")
    newfile.write(encoded_message)
    newfile.close()
    file.close()

#Pdf Decoder function
def decoder_pdf(filename, saving_filename00):
    if saving_filename00.endswith('.txt') == False:
        saving_filename01 = saving_filename00 + '.txt'
    else:
         saving_filename01 = saving_filename00
    pdf = pdfplumber.open(filename)
    pdf_text = ""
    for page in pdf.pages:
        if page.extract_text() == None:
            pdf_text += "No text found for this page" + "\n-------Next page------\n"
        else:
            pdf_text += page.extract_text() + "\n-------Next page------\n"
    pdf.close()    
    ux1 = list(pdf_text)
    for n in range(len(ux1)):
        something = ux1[n]
        ux1[n] = reverse_dict0.get(something, something)
    encoded_message = ''.join(ux1)
    savingfile = open(saving_filename01, "w")
    savingfile.write(encoded_message)
    savingfile.close()
    print(f"\nCongrats!! File saved as {saving_filename01}\n")
    
    
# docx deocder func.


def decoder_docx(filename, saving_filename00):
    if filename.endswith(".docx") == False:
        finalfilename = filename + ".docx"
    else:
        finalfilename = filename
    
    
    if saving_filename00.endswith('.txt') == False:
        saving_filename01 = saving_filename00 + '.txt'
    else:
         saving_filename01 = saving_filename00
         
    doc = docx.Document(finalfilename)
    doc_text = ""
    for para in doc.paragraphs:
        if para.text == "":
            doc_text += "No text found for this para" + "\n-------Next para------\n"
        else:
            doc_text += para.text + "\n-------Next page------\n"
    
    ux1 = list(doc_text)
    for n in range(len(ux1)):
        something = ux1[n]
        ux1[n] = reverse_dict0.get(something, something)
    encoded_message = ''.join(ux1)
    savingfile = open(saving_filename01, "w", encoding="utf-8")
    savingfile.write(encoded_message)
    savingfile.close()
    print(f"\nCongrats!! File saved as {saving_filename01}\n")


#Final execution
while True:
    option_for_filetype = input("""So what kind of doc. you want to encode or decode ? 
    Kindly choose one of the following option -
    a. Plain Text
    b. \".txt\" file
    c. \".pdf\" file
    d. \".docx\" file 
    Choose one - a/b/c/d :""").upper()
    
    if option_for_filetype == "A":
        message = input("Kindly enter your message : ")
        while True:
            enorden_fora = input("""So what do you want to do ?
            a. Encode
            b. Decode
            c. Exit
            Kindly a option : """).lower() 
            
            if enorden_fora == "a":
                encoder_usual(message)
                break
                
            elif enorden_fora == "b":
                decoder_usual(message)
                break
            
            elif enorden_fora == "c":
                print("Getting back..............\n\n")
                break
            
            else:
                print("Invalid input, kindly enter - a/b/c - \n")
                
    elif option_for_filetype == "B":
        enorden_forb = input("""What do you want to do -
        a. Encode
        b. Decode
        Kindly choose a/b : """).lower()
        while True:
             print("""So, kindly follow these instructions -
                1. Just put the file name if the file is in same directory -
                      For e.g. just "Filename.txt" or "Filename"
                2. Put absolute path of the file if its not in the same directory -
                      For e.g. /storage/emulated/0/filename.txt (absolute)""")
             if enorden_forb == "a":
                 filename_final_00 = input("\n\nKindly enter th file name or its absolute path : ") 
                 saving_filename_temp = input("\nEnter file name, by which you want to save it : ")
                 if saving_filename_temp.endswith(".txt") == False:
                     saving_final_filename = saving_filename_temp + ".txt"
                 else:
                     saving_final_filename = saving_filename_temp
                 encoder_txt(filename_final_00, saving_final_filename)
                 print("\n\nCongrats, file successfully saved !!\n\n")
                 break
               
             elif enorden_forb == "b":
                 filename_final_00 = input("\n\nKindly enter th file name or its absolute path : ") 
                 saving_filename_temp = input("\nEnter file name, by which you want to save it : ")
                 if saving_filename_temp.endswith(".txt") == False:
                     saving_final_filename = saving_filename_temp + ".txt"
                 else:
                     saving_final_filename = saving_filename_temp
                 decoder_txt(filename_final_00, saving_final_filename)
                 print("\n\nCongrats, file successfully saved !!\n\n")
                 break
             else:
                 print("\nIt's a invalid input, kindly enter a or b ")
        
              
            
    elif option_for_filetype == "C":
        enorden_forc = input("""What do you want to do -
        a. Encode
        b. Decode
        Kindly choose a/b : """).lower()
        while True:
             print("""So, kindly follow these instructions -
                1. Just put the file name if the file is in same directory -
                      For e.g. just "Filename.pdf" or "Filename"
                2. Put absolute path of the file if its not in the same directory -
                      For e.g. /storage/emulated/0/filename.pdf (absolute)\n""")
             if enorden_forc == 'a':
                  filename = input("Kindly enter filename/absolute path : ")
                  if filename.endswith(".pdf") == False:
                      finalfilename = filename + ".pdf"
                  else:
                      finalfilename = filename
                  savingname001 = input("From which name do you want to save encoded text : ")    
                  if savingname001.endswith(".txt") == False:
                     savingname002  = savingname001  + ".txt"
                  else:
                     savingname002  = savingname001
                  encoder_pdf(finalfilename, savingname002)
                  break
             elif enorden_forc == 'b':
                 filename = input("Kindly enter filename/absolute path : ")
                 if filename.endswith(".pdf") == False:
                     finalfilename = filename + ".pdf"
                 else:
                     finalfilename = filename
                 savingname001 = input("From which name do you want to save encoded text : ")                             
                 if savingname001.endswith(".txt") == False:
                     savingname002  = savingname001  + ".txt"
                 else:
                     savingname002  = savingname001 
                 decoder_pdf(finalfilename, savingname002)
                 break
             else:
                  print('invalid input, kindly enter a/b -\n\n')
                  
     
                  
    elif option_for_filetype == "D":
        enorden_ford = input("""What do you want to do -
        a. Encode
        b. Decode
        Kindly choose a/b : """).lower()
        while True:
             print("""So, kindly follow these instructions -
                1. Just put the file name if the file is in same directory -
                      For e.g. just "Filename.docx" or "Filename"
                2. Put absolute path of the file if its not in the same directory -
                      For e.g. /storage/emulated/0/filename.docx (absolute)\n""")
             if enorden_ford == "a":
                  filename = input("Kindly enter the name of file you want to encode : \n")
                  savingfilename = input("Kindly enter the name of the file you want to save with : \n")
                  encoder_docx(filename, savingfilename)
                  print(f"Congrats !! , text of {filename} has been encoded and saved as {savingfilename  }")
                  
             elif enorden_ford == "b":
                  filename = input("Kindly enter the name of file you want to decode : \n")
                  savingfilename = input("Kindly enter the name of the file you want to save with : \n")
                  decoder_docx(filename, savingfilename)
                  print(f"Congrats !! , text of {filename} has been decoded and saved as {savingfilename  }")
             else:
                  print("Incorrect input, kindly enter a/b and according to instruction")
    else:
        print("Invalid input, kindly choose options - a/b/c/d \n")
    