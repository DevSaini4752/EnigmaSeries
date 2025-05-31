"""- Not going to rewrite the whole code
- Will use decrypt_plain.py for decryption of text
- Will save a lot of time
- Will take out text from Word file
- Will put it into decrypt_plain.py module, and it's func
- Will save the text in the Word file
"""

# Importing modules
import decrypt_plain  # File for decryption
import docx  # Module for handling word


# Making the function for the process.
def decryption_doc(filename_doc, saving_name_doc, decrypt_token):
    """- First, Going to handle correction
        - Check whether file name ende with proper extension (.docx) or not
        - If not then will add
        - If there then pass
        - Then the code for taking out text from file
        - Then will put the text into decryption_plain function
        - Will make a file with the same format
        - Will save the text in it
        - Not going to handle error here, have a special module for it"""

    # Correction handling, correcting the user input
    if not filename_doc.endswith(".docx"):
        # Adding .docx to filename_doc if it's not present
        filename_doc += ".docx"

    if not saving_name_doc.endswith(".docx"):
        # Adding .docx to saving file name if it's not there
        saving_name_doc += ".docx"

    # Opening the file and taking text
    file_out = docx.Document(filename_doc)  # Opening word
    full_text = []

    # Taking text and putting the text into a str
    for para in file_out.paragraphs:
        full_text.append(para.text)
    full_text_final = "\n".join(full_text)

    # decrypting the message
    decrypted_msg = decrypt_plain.decryption_plain(decrypt_token, full_text_final)


    # Saving the msg in a Word file
    saving_file = docx.Document()
    saving_file.add_paragraph(decrypted_msg)
    saving_file.save(saving_name_doc)

    # Telling user the file saved successfully
    print("File saved successfully!!!")



if __name__ == "__main__":
    print("Testing the method")
    trial01 = input("Kindly enter the word file name : ")
    trial02 = input("Kindly enter the name by which you want to save it : ")
    code = input("Kindly enter the decryption token : ")
    decryption_doc(trial01, trial02, code)