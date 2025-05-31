"""- Not going to rewrite the whole code
- Will use encrypt_plain.py for encryption of text
- Will take out text from Word file
- Will put it into encrypt_plain.py module, and it's function
- Will save the text in Word file
- Will save a lot of time
"""


#Importing modules
import encrypt_plain # File for encryption
import docx # Module for handling word

# Making the function for process
# It will return decryption token
def encryption_doc(filename_doc, saving_name_doc):
    """- First, Going to handle correction
        - Check whether file name ende with proper extension (.docx) or not
        - If not then will add
        - If there then pass
        - Then the code for taking out text from file
        - Then will put the text into encryption_plain function
        - Will make a file with the same format
        - Will save the text in it
        - Not going to handle error here, have a special module for it"""


    #Correction handling, correcting the user input
    if not filename_doc.endswith(".docx"):
        # Adding .docx to filename_doc if it's not present
        filename_doc += ".docx"

    if not saving_name_doc.endswith(".docx"):
        # Adding .docx to saving file name if it's not there
        saving_name_doc += ".docx"

    #Opening the file and taking text
    file_in = docx.Document(filename_doc) #Opening word
    full_text = []
    
    #Taking text and putting the text into a str
    for para in file_in.paragraphs:
        full_text.append(para.text)
    full_text_final = "\n".join(full_text)

    #Encrypting the message
    token_msg_doc = encrypt_plain.encryption_plain(full_text_final)
    token_doc = token_msg_doc[0]
    encrypted_msg = token_msg_doc[1]

    #Saving the msg in a Word file
    saving_file = docx.Document()
    saving_file.add_paragraph(encrypted_msg)
    saving_file.save(saving_name_doc)

    
    #Telling user the file saved successfully
    print("File saved successfully!!!")

    #Returning the token as a message is already saved to the file
    return token_doc

if __name__ == "__main__":
    print("Testing the method")
    trial01 = input("Kindly enter the word file name : ")
    trial02 = input("Kindly enter the name by which you want to save it : ")
    print(encryption_doc(trial01, trial02))