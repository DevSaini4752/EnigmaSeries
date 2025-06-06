#Importing modules
import pdfplumber
#Kindly remember that this docx module is not the old one, so if you are installing module then kindly install "python-docx"
#this one is new, better and correct version which our project will expect here
import docx
import random

'''So here in my code project is cipher, a dynamic cipher. It takes the message in the form of "plain text", ".txt", ".pdf", and ".docx". 
So this script takes the message of the person and encrypt that message with a special key which is only for that message, so even if the 
3rd person tries to open it from my program then it woundn't be able to, because of lacking key. And predicting the key is almost imposs-
ible, because the key is of 90 digits, which makes 10^80 cobmination. As a total i made 4 cipher, first was very simple, in its logic it 
only added random character in between letters and could be easily decoded. Then in second cipher i upgraded it a little bit, and logic 
was much better, the messanger and receiver had a key which decodes their message but if the key is leaked then it was over.'''
#Assigning letters
# Making the function for making a new key every time so
#This is the debugged, upgraded and correct version of Enigma 2.0 - 11:27 a.m. - 11-04-2025
def dynamic_key():
    global list_keys
    list_keys = ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'J', 'K', 'L', 'M', 'N', 'O', 'P', 'Q', 'R', 'S', 'T',
                 'U', 'V', 'W', 'X', 'Y', 'Z', 'a', 'b', 'c', 'd', 'e', 'f', 'g', 'h', 'i', 'j', 'k', 'l', 'm', 'n',
                 'o', 'p', 'q', 'r', 's', 't', 'u', 'v', 'w', 'x', 'y', 'z', '0', '1', '2', '3', '4', '5', '6', '7',
                 '8', '9', ' ', '#', '.', ',', '!', '?', '-', '_', '(', ')', '[', ']', '{', '}', '<', '>', '@', '$',
                 '%', '^', '&', '*', '+', '=', '/', '\\', '|', ';', ':', "'", '"', '`']
    global list_values
    list_values = ['M', 'Q', 'W', 'E', 'R', 'T', 'Y', 'U', 'I', 'O', 'P', 'A', 'S', 'D', 'F', 'G', 'H', 'J', 'K', 'L',
                   'Z', 'X', 'C', 'V', 'B', 'N', 'm', 'q', 'w', 'e', 'r', 't', 'y', 'u', 'i', 'o', 'p', 'a', 's', 'd',
                   'f', 'g', 'h', 'j', 'k', 'l', 'z', 'x', 'c', 'v', 'b', 'n', '9', '8', '7', '6', '5', '4', '3', '2',
                   '1', '0', '#', ' ', ',', '.', '?', '!', '_', '-', ')', '(', ']', '[', '}', '{', '>', '<', '$', '@',
                   '^', '%', '*', '&', '=', '+', '\\', '/', ';', '|', "'", ':', '`', '"']
    random.shuffle(list_values)
    global dict0
    dict0 = dict(zip(list_keys,list_values))
    code = ''.join(dict0.values())
    print(f"So here is ur key for the code - {code}")
    return code


def code_to_dictionary(code):
    listvalues_for_code = list(code)
    global reverse_dict0
    reverse_dict0 = dict(zip(listvalues_for_code, list_keys))

code_to_dictionary(dynamic_key())
print("Kindly ignore the above text ---")
# Usual Encoder func.

def encoder_usual(x00):
    dynamic_key()
    ux1 = list(x00)
    for n in range(len(ux1)):
        something = ux1[n]
        ux1[n] = dict0.get(something, something)
    encoded_message = ''.join(ux1)
    print(f"Your encoded message is - {encoded_message}\n")
    
# Txt Encoder Function

def encoder_txt(filename, savingfilename):
    dynamic_key()
    if filename.endswith(".txt") == False:
        finalfile = filename + ".txt"
    else:
        finalfile = filename

    file = open(finalfile, "r")
    intext = file.read()
    ux1 = list(intext)
    for n in range(len(ux1)):
        something = ux1[n]
        ux1[n] = dict0.get(something, something)

    encoded_message = ''.join(ux1)
    newfile = open(savingfilename, "w")
    newfile.write(encoded_message)
    newfile.close()
    file.close()
    
    
#Pdf encoding function
def encoder_pdf(filename, saving_filename00):
    dynamic_key()
    if saving_filename00.endswith('.txt') == False:
        saving_filename01 = saving_filename00 + '.txt'
    else:
         saving_filename01 = saving_filename00
    pdf = pdfplumber.open(filename)
    pdf_text = ""
    for page in pdf.pages:
        if page.extract_text() == None:
            pdf_text += "No text found for this page" + "\n-------Next page------\n"
        else:
            pdf_text += page.extract_text() + "\n-------Next page------\n"
    pdf.close()
    ux1 = list(pdf_text)
    for n in range(len(ux1)):
        something = ux1[n]
        ux1[n] = dict0.get(something, something)
    encoded_message = ''.join(ux1)
    savingfile = open(saving_filename01, "w")
    savingfile.write(encoded_message)
    savingfile.close()
    print(f"\nCongrats!! File saved as {saving_filename01}\n")

    
#Docx encoder function :-)

def encoder_docx(filename, saving_filename00):
    dynamic_key()
    if filename.endswith(".docx") == False:
        finalfilename = filename + ".docx"
    else:
        finalfilename = filename

    
    if saving_filename00.endswith('.txt') == False:
        saving_filename01 = saving_filename00 + '.txt'
    else:
        saving_filename01 = saving_filename00
         
    doc = docx.Document(finalfilename)
    doc_text = ""
    for para in doc.paragraphs:
        if para.text == "":
            doc_text += "No text found for this para" + "\n-------Next para------\n"
        else:
            doc_text += para.text + "\n-------Next page------\n"
   
    ux1 = list(doc_text)
    for n in range(len(ux1)):
        something = ux1[n]
        ux1[n] = dict0.get(something, something)
    encoded_message = ''.join(ux1)
    savingfile_encoder_func = open(saving_filename01, "w", encoding="utf-8")
    savingfile_encoder_func.write(encoded_message)
    savingfile_encoder_func.close()
    print(f"\nCongrats!! File saved as {saving_filename01}\n") 
                
#Decoder func.

def decoder_usual(filename):
    code = input("Kindly enter the code, otherwise you will get a wrong message : ")
    code_to_dictionary(code)
    ux1 = list(filename)
    for n in range(len(ux1)):
        something = ux1[n]
        ux1[n] = reverse_dict0.get(something, something)
    decoded_message = ''.join(ux1)
    print("Your decoded message is - ", decoded_message, "\n")
    
#Txt Decoder function

def decoder_txt(filename, savingfilename):
    code = input("Kindly enter the code, otherwise you will get a wrong message : ")
    code_to_dictionary(code)
    if filename.endswith(".txt") == False:
        finalfile = filename + ".txt"
    else:
        finalfile = filename
        
    file = open(finalfile, "r") 
    intext = file.read()
    ux1 = list(intext)
    for n in range(len(ux1)):
        something = ux1[n]
        ux1[n] = reverse_dict0.get(something, something)
        
    encoded_message = ''.join(ux1)
    newfile = open(savingfilename, "w")
    newfile.write(encoded_message)
    newfile.close()
    file.close()

#Pdf Decoder function
def decoder_pdf(filename, saving_filename00):
    code = input("Kindly enter the code, otherwise you will get a wrong message : ")
    code_to_dictionary(code)
    if saving_filename00.endswith('.txt') == False:
        saving_filename01 = saving_filename00 + '.txt'
    else:
         saving_filename01 = saving_filename00
    pdf = pdfplumber.open(filename)
    pdf_text = ""
    for page in pdf.pages:
        if page.extract_text() == None:
            pdf_text += "No text found for this page" + "\n-------Next page------\n"
        else:
            pdf_text += page.extract_text() + "\n-------Next page------\n"
    pdf.close()    
    ux1 = list(pdf_text)
    for n in range(len(ux1)):
        something = ux1[n]
        ux1[n] = reverse_dict0.get(something, something)
    encoded_message = ''.join(ux1)
    savingfile = open(saving_filename01, "w")
    savingfile.write(encoded_message)
    savingfile.close()
    print(f"\nCongrats!! File saved as {saving_filename01}\n")
    
    
# docx deocder func.


def decoder_docx(filename, saving_filename00):
    code = input("Kindly enter the code, otherwise you will get a wrong message : ")
    code_to_dictionary(code)
    if filename.endswith(".docx") == False:
        finalfilename = filename + ".docx"
    else:
        finalfilename = filename

    if saving_filename00.endswith('.txt') == False:
        saving_filename01 = saving_filename00 + '.txt'
    else:
         saving_filename01 = saving_filename00
         
    doc = docx.Document(finalfilename)
    doc_text = ""
    for para in doc.paragraphs:
        if para.text == "":
            doc_text += "No text found for this para" + "\n-------Next para------\n"
        else:
            doc_text += para.text + "\n-------Next page------\n"
    
    ux1 = list(doc_text)
    for n in range(len(ux1)):
        something = ux1[n]
        ux1[n] = reverse_dict0.get(something, something)
    encoded_message = ''.join(ux1)
    savingfile = open(saving_filename01, "w", encoding="utf-8")
    savingfile.write(encoded_message)
    savingfile.close()
    print(f"\nCongrats!! File saved as {saving_filename01}\n")


#Final execution
while True:


    option_for_filetype = input("""So what kind of doc. you want to encode or decode ? 
    Kindly choose one of the following option -
    a. Plain Text
    b. \".txt\" file
    c. \".pdf\" file
    d. \".docx\" file 
    Choose one - a/b/c/d :""").upper()
    
    if option_for_filetype == "A":
        message = input("Kindly enter your message : ")
        while True:
            try:
                enorden_fora = input("""So what do you want to do ?
                a. Encode
                b. Decode
                c. Exit
                Kindly a option : """).lower()

                if enorden_fora == "a":
                    encoder_usual(message)
                    break

                elif enorden_fora == "b":
                    decoder_usual(message)
                    break

                elif enorden_fora == "c":
                    print("Getting back..............\n")
                    break

                else:
                    print("Invalid input, kindly enter - a/b/c - \n")
                break
            except Exception as ex:
                print(f"Error - {ex}")


    elif option_for_filetype == "B":
        enorden_forb = input("""What do you want to do -
        a. Encode
        b. Decode
        Kindly choose a/b : """).lower()
        while True:
            try:
                print("""So, kindly follow these instructions -
                   1. Just put the file name if the file is in same directory -
                      For e.g. just "Filename.txt" or "Filename"
                   2. Put absolute path of the file if its not in the same directory -
                      For e.g. /storage/emulated/0/filename.txt (absolute)""")
                if enorden_forb == "a":
                    filename_final_00 = input("\n\nKindly enter th file name or its absolute path : ")
                    saving_filename_temp = input("\nEnter file name, by which you want to save it : ")
                    if saving_filename_temp.endswith(".txt") == False:
                        saving_final_filename = saving_filename_temp + ".txt"
                    else:
                        saving_final_filename = saving_filename_temp
                    encoder_txt(filename_final_00, saving_final_filename)
                    print("\n\nCongrats, file successfully saved !!\n\n")
                    break

                elif enorden_forb == "b":
                    filename_final_00 = input("\n\nKindly enter th file name or its absolute path : ")
                    saving_filename_temp = input("\nEnter file name, by which you want to save it : ")
                    if saving_filename_temp.endswith(".txt") == False:
                        saving_final_filename = saving_filename_temp + ".txt"
                    else:
                        saving_final_filename = saving_filename_temp
                    decoder_txt(filename_final_00, saving_final_filename)
                    print("\n\nCongrats, file successfully saved !!\n")
                    break
                else:
                    print("\nIt's a invalid input, kindly enter a or b ")
            except Exception as ex:
                print(f"Error - {ex}")

              
            
    elif option_for_filetype == "C":
        enorden_forc = input("""What do you want to do -
        a. Encode
        b. Decode
        Kindly choose a/b : """).lower()
        while True:
            try:
                print("""So, kindly follow these instructions -
                1. Just put the file name if the file is in same directory -
                        For e.g. just "Filename.pdf" or "Filename"
                2. Put absolute path of the file if its not in the same directory -
                        For e.g. /storage/emulated/0/filename.pdf (absolute)\n""")
                if enorden_forc == 'a':
                    filename = input("Kindly enter filename/absolute path : ")
                    if filename.endswith(".pdf") == False:
                        finalfilename = filename + ".pdf"
                    else:
                        finalfilename = filename
                    savingname001 = input("From which name do you want to save encoded text : ")
                    if savingname001.endswith(".txt") == False:
                     savingname002  = savingname001  + ".txt"
                    else:
                        savingname002  = savingname001
                    encoder_pdf(finalfilename, savingname002)
                    break
                elif enorden_forc == 'b':
                    filename = input("Kindly enter filename/absolute path : ")
                    if filename.endswith(".pdf") == False:
                        finalfilename = filename + ".pdf"
                    else:
                        finalfilename = filename
                    savingname001 = input("From which name do you want to save encoded text : ")
                    if savingname001.endswith(".txt") == False:
                        savingname002  = savingname001  + ".txt"
                    else:
                        savingname002  = savingname001
                    decoder_pdf(finalfilename, savingname002)
                    break
                else:
                    print('invalid input, kindly enter a/b -\n')
            except Exception as ex:
                print(f"Error - {ex}")
                  
     
                  
    elif option_for_filetype == "D":
        while True:
            enorden_ford = input("""What do you want to do -
            a. Encode
            b. Decode
            c. Exit
            Kindly choose a/b : """).lower()

            try:
                print("""So, kindly follow these instructions -
                1. Just put the file name if the file is in same directory -
                        For e.g. just "Filename.docx" or "Filename"
                2. Put absolute path of the file if its not in the same directory -
                        For e.g. /storage/emulated/0/filename.docx (absolute)\n""")
                if enorden_ford == "a":
                    filename = input("Kindly enter the name of file you want to encode : \n")
                    savingfilename = input("Kindly enter the name of the file you want to save with : \n")
                    encoder_docx(filename, savingfilename)
                    print(f"Congrats !! , text of {filename} has been encoded and saved as {savingfilename  }")

                elif enorden_ford == "b":
                    filename = input("Kindly enter the name of file you want to decode : \n")
                    savingfilename = input("Kindly enter the name of the file you want to save with : \n")
                    decoder_docx(filename, savingfilename)
                    print(f"Congrats !! , text of {filename} has been decoded and saved as {savingfilename  }")
                elif enorden_ford == "c":
                    print("Getting back............/")
                    break
                else:
                    print("Incorrect input, kindly enter a/b and according to instruction")
            except Exception as ex:
                print(f"Error - {ex}")
    else:
        print("Invalid input, kindly choose options - a/b/c/d \n")
    